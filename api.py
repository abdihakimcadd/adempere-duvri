from fastapi import FastAPI, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from supabase import create_client, Client
import os
import io
import json
from datetime import datetime

# DOCX Imports ONLY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import your pipeline functions
from knowledge import get_context_block
from agent_drafter import draft_duvri
from agent_reviewer import review_duvri

app = FastAPI()

# UPDATED CORS TO FIX THE NETWORK ERROR (Added Lovable preview URL)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://adempere.com", 
        "https://www.adempere.com",
        "https://easy-deal-vista.lovable.app",  # Lovable Preview URL
        "http://localhost:3000",
        "http://localhost:5173"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Supabase Client
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

def generate_duvri_docx(data):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # === HEADER ===
    title = doc.add_heading('DOCUMENTO UNICO DI VALUTAZIONE DEI RISCHI DA INTERFERENZA (DUVRI)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    subtitle = doc.add_paragraph('Redatto ai sensi dell\'Art. 26, comma 3, D.Lgs. 9 aprile 2008 n. 81')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # === PARTE 1: COMMITTENTE ===
    doc.add_heading('Parte 1 — Azienda Committente', level=1)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    cells1 = [
        ('Ragione Sociale:', data['host_company']),
        ('P.IVA / C.F.:', data['host_vat']),
        ('Sede dei Lavori:', data['work_site']),
        ('Ruolo Redattore:', 'RSPP / Datore di Lavoro Committente')
    ]
    for i, (label, value) in enumerate(cells1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
        table1.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # === PARTE 5: APPALTATRICE ===
    doc.add_heading('Parte 5 — Azienda Appaltatrice', level=1)
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    cells2 = [
        ('Ragione Sociale:', data['contractor_name']),
        ('P.IVA / C.F.:', data['contractor_vat']),
        ('Oggetto Appalto:', data['work_type']),
        ('Periodo Lavori:', f"Da {data['start_date']} a {data['end_date']}")
    ]
    for i, (label, value) in enumerate(cells2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value
        table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # === PARTE 4.A: RISK MATRIX ===
    doc.add_heading('Parte 4.A — Matrice dei Rischi Interferenti', level=1)
    risk_table = doc.add_table(rows=1 + len(data['risks']), cols=4)
    risk_table.style = 'Table Grid'
    hdr_cells = risk_table.rows[0].cells
    headers = ['Rischio Identificato', 'Probabilità (P)', 'Magnitudo (M)', 'Indice (P × M)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for i, risk in enumerate(data['risks']):
        row = risk_table.rows[i+1].cells
        row[0].text = risk['name']
        row[1].text = str(risk['probability'])
        row[2].text = str(risk['magnitude'])
        row[3].text = f"{risk['score']} ({risk['level']})"
    
    doc.add_paragraph('Classificazione: 1-4 Basso, 5-9 Medio, 10-16 Alto.').runs[0].font.italic = True
    
    doc.add_page_break()
    
    # === PARTE 2: VALUTAZIONE RICOGNITIVA ===
    doc.add_heading('Parte 2 — Valutazione Ricognitiva (Art. 26, comma 3-ter)', level=1)
    doc.add_paragraph(data['valutazione_text'])
    
    doc.add_page_break()
    
    # === PARTE 4.B: MISURE ===
    doc.add_heading('Parte 4.B — Misure di Prevenzione e Protezione', level=1)
    doc.add_paragraph(data['misure_text'])
    
    doc.add_page_break()
    
    # === COSTI ===
    doc.add_heading('Individuazione dei Costi della Sicurezza (Art. 26, comma 5)', level=1)
    doc.add_paragraph(data['costi_text'])
    
    doc.add_page_break()
    
    # === PARTE 3: REGOLE ===
    doc.add_heading('Parte 3 — Regole di Prevenzione ed Emergenza in vigore nel Sito', level=1)
    rules = [
        'Divieto di fumo in tutti i locali aziendali.',
        'Obbligo di utilizzo dei DPI ove previsti dalla segnaletica di sicurezza.',
        'Mantenimento delle vie di fuga e dei presidi antincendio sgombri da ostacoli.',
        'Rispetto delle procedure di emergenza aziendali (vie di fuga, punti di raccolta).'
    ]
    for rule in rules:
        p = doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # === PARTE 4.C: COORDINAMENTO ===
    doc.add_heading('Parte 4.C — Procedure di Coordinamento', level=1)
    doc.add_paragraph(data['coordination_text'])
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.rows[0].cells[0].text = '___________________________'
    sig_table.rows[0].cells[1].text = '___________________________'
    sig_table.rows[1].cells[0].text = 'Il Committente'
    sig_table.rows[1].cells[1].text = "L'Appaltatore"
    sig_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    sig_table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
    sig_table.rows[2].cells[0].text = data['host_company']
    sig_table.rows[2].cells[1].text = data['contractor_name']
    
    doc.add_page_break()
    
    # === ALLEGATO A ===
    doc.add_heading('Allegato A — Estratto Normativo (Art. 26, D.Lgs 81/08)', level=1)
    
    legal_articles = [
        ("Art. 26, comma 1:", "Il datore di lavoro committente verifica l'idoneità tecnico-professionale delle imprese appaltatrici e fornisce dettagliate informazioni sui rischi specifici esistenti nell'ambiente e sulle misure di prevenzione e di emergenza adottate."),
        ("Art. 26, comma 3:", "Il datore di lavoro committente elabora un unico documento di valutazione dei rischi che indichi le misure adottate per eliminare o, ove ciò non è possibile, ridurre al minimo i rischi da interferenze. Il documento è allegato al contratto di appalto o di opera e deve essere adeguato in funzione dell'evoluzione dei lavori, servizi e forniture."),
        ("Art. 26, comma 3-ter:", "Il soggetto che affida il contratto redige il documento di valutazione dei rischi da interferenze recante una valutazione ricognitiva dei rischi standard relativi alla tipologia della prestazione."),
        ("Art. 26, comma 5:", "Nei contratti devono essere specificamente indicati i costi delle misure adottate per eliminare o ridurre al minimo i rischi da interferenze. Tali costi non sono soggetti a ribasso."),
        ("Art. 55, comma 5, lett. d):", "Sanzione per mancata redazione del DUVRI: Arresto da due a quattro mesi o ammenda da € 1.500 a € 6.000.")
    ]
    
    for title, text in legal_articles:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(' ' + text)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Documento generato da Adempere — AI Compliance Platform. Questo DUVRI deve essere allegato al contratto di appalto e aggiornato in funzione dell\'evoluzione dei lavori (Art. 26, comma 3, D.Lgs 81/08).')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@app.post("/api/start-duvri")
async def start_duvri_pipeline(payload: dict, background_tasks: BackgroundTasks):
    try:
        risks_input = payload.get('risks', [])
        work_type = payload.get('work_type', 'General Maintenance')
        
        risk_translation = {
            "work at heights": "Lavori in quota",
            "electrocution": "Rischio Elettrico (Folgorazione)",
            "noise": "Rumore",
            "chemical agents": "Agenti Chimici",
            "confined spaces": "Spazi Confinati",
            "manual handling": "Movimentazione Manuale dei Carichi"
        }
        
        risk_scores = {}
        for risk in risks_input:
            risk_it = risk_translation.get(risk.lower(), risk.title())
            if 'quota' in risk_it.lower() or 'height' in risk.lower():
                p, m = 3, 4
            elif 'chimic' in risk_it.lower() or 'chemical' in risk.lower():
                p, m = 2, 3
            elif 'rumore' in risk_it.lower() or 'noise' in risk.lower():
                p, m = 4, 2
            elif 'elettric' in risk_it.lower() or 'electric' in risk.lower():
                p, m = 2, 4
            else:
                p, m = 3, 3
            
            score = p * m
            risk_scores[risk_it] = {"P": p, "M": m, "Score": score}
            
    except Exception as e:
        return {"error": f"Invalid input format: {str(e)}"}
    
    today_date = datetime.now().strftime("%d/%m/%Y")
    payload['risks_it'] = list(risk_scores.keys())
    payload['risk_scores'] = risk_scores
    payload['today_date'] = today_date
    
    # SAVE TO SUPABASE DB & START AGENTS (Using 'documents' table)
    response = supabase.table("documents").insert({
        "user_id": payload.get("user_id"),
        "status": "PROCESSING",
        "email": payload.get("email", "unknown@example.com"),
        "document_type": "DUVRI",
        "intake_answers": payload
    }).execute()
    job_id = response.data[0]['id']
    
    background_tasks.add_task(run_duvri_task, job_id, payload)
    return {"job_id": job_id, "status": "PROCESSING"}

@app.get("/api/status/{job_id}")
def get_status(job_id: str):
    # Using 'documents' table
    response = supabase.table("documents").select("status, pdf_url, error").eq("id", job_id).execute()
    if response.data:
        return response.data[0]
    return {"error": "Job not found"}

def run_duvri_task(job_id: str, frontend_json: dict):
    try:
        print(f"[{job_id}] Starting DUVRI Pipeline...")
        context_block = get_context_block()
        draft_json = draft_duvri(frontend_json, context_block)
        
        # AI PROOFREADING
        draft_string = json.dumps(draft_json, ensure_ascii=False)
        draft_string = draft_string.replace("c. c.", "c.c.")
        draft_string = draft_string.replace(". .", ".")
        draft_string = draft_string.replace("..", ".")
        draft_string = draft_string.replace("  ", " ")
        draft_json = json.loads(draft_string)
        
        # AGENT REVIEWER
        print(f"[{job_id}] Reviewing DUVRI for compliance...")
        corrected_json = review_duvri(draft_json, frontend_json, context_block)
        corrected_json.pop("review_notes", None)
        draft_json = corrected_json

        # MAP DATA FOR DOCX (Safe .get() fallbacks to prevent KeyError crashes)
        risks_list = []
        for risk_name, scores in frontend_json.get('risk_scores', {}).items():
            score = scores['Score']
            if score <= 4:
                level = "Basso"
            elif score <= 9:
                level = "Medio"
            else:
                level = "Alto"
                
            risks_list.append({
                'name': risk_name,
                'probability': scores['P'],
                'magnitude': scores['M'],
                'score': score,
                'level': level
            })
            
        pdf_data = {
            'host_company': frontend_json.get('host_company', 'N/D'),
            'host_vat': frontend_json.get('host_vat', 'N/D'),
            'work_site': frontend_json.get('work_site_address', 'N/D'),
            'contractor_name': frontend_json.get('contractor_name', 'N/D'),
            'contractor_vat': frontend_json.get('contractor_vat', 'N/D'),
            'work_type': frontend_json.get('work_type', 'N/D'),
            'start_date': frontend_json.get('start_date', 'N/D'),
            'end_date': frontend_json.get('end_date', 'N/D'),
            'valutazione_text': draft_json.get('valutazione_ricognitiva', ''),
            'misure_text': draft_json.get('risk_measures', ''),
            'costi_text': draft_json.get('safety_costs', ''),
            'coordination_text': draft_json.get('coordination_procedures', ''),
            'risks': risks_list
        }
        
        # GENERATE DOCX
        try:
            print(f"[{job_id}] Converting to DOCX...")
            docx_bytes = generate_duvri_docx(pdf_data)
        except Exception as docx_err:
            print(f"[{job_id}] DOCX Generation Failed: {str(docx_err)}")
            raise docx_err
        
        file_name = f"DUVRI_{pdf_data['contractor_name'].replace(' ', '_')}_{job_id}.docx"
        print(f"[{job_id}] Uploading to Supabase Storage...")
        supabase.storage.from_("dvr-documents").upload(
            file_name, 
            docx_bytes, 
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )
        
        public_url = supabase.storage.from_("dvr-documents").get_public_url(file_name)
        
        print(f"[{job_id}] Updating DB with DOCX URL...")
        # Using 'documents' table
        supabase.table("documents").update({
            "status": "COMPLETED",
            "pdf_url": public_url  # Keeping column name pdf_url for frontend compatibility
        }).eq("id", job_id).execute()
        
        print(f"[{job_id}] DUVRI Complete!")
        
    except Exception as e:
        print(f"[{job_id}] Error: {str(e)}")
        # Using 'documents' table
        supabase.table("documents").update({
            "status": "FAILED",
            "error": str(e)
        }).eq("id", job_id).execute()
